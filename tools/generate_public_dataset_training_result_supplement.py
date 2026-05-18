import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
ART_DIR = ROOT / "backend" / "models" / "thyroid_resnet50_tn5000_formal_20260507_024833_artifacts"
OUT_PATH = Path(r"C:\Users\86136\Downloads\公开数据集训练结果补充说明-正文版.docx")


def zh(text_ascii_escaped: str) -> str:
    return text_ascii_escaped.encode("ascii").decode("unicode_escape")


def load_summary(path: Path) -> dict:
    summary = {}
    for line in path.read_text(encoding="utf-8").splitlines():
        if "=" in line:
            k, v = line.split("=", 1)
            summary[k.strip()] = v.strip()
    return summary


def set_run_font(run, font_name="宋体", size=12, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_format(paragraph, align=None, first_indent=True, line_spacing=1.5):
    fmt = paragraph.paragraph_format
    fmt.line_spacing = line_spacing
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if first_indent:
        fmt.first_line_indent = Pt(24)
    if align is not None:
        paragraph.alignment = align


def add_text_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    set_paragraph_format(p)
    r = p.add_run(text)
    set_run_font(r, "宋体", 12, False)
    return p


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, line_spacing=1.25)
    r = p.add_run(text)
    set_run_font(r, "宋体", 11, False)
    return p


def set_table_text(cell, text: str, bold=False, align="center"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, "宋体", 10.5, bold)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = f"w:{edge}"
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ("val", "sz", "space", "color"):
                if key in edge_data:
                    element.set(qn("w:" + key), str(edge_data[key]))


def style_table(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": 8, "color": "000000"},
                bottom={"val": "single", "sz": 8, "color": "000000"},
                left={"val": "single", "sz": 8, "color": "000000"},
                right={"val": "single", "sz": 8, "color": "000000"},
            )


def main():
    summary = load_summary(ART_DIR / "summary.txt")
    best_val = json.loads((ART_DIR / "best_val_report.json").read_text(encoding="utf-8"))
    test_report = json.loads((ART_DIR / "test_report.json").read_text(encoding="utf-8"))
    rows = list(csv.DictReader((ART_DIR / "metrics.csv").open("r", encoding="utf-8-sig")))
    best_row = max(rows, key=lambda r: float(r["val_acc"]))

    train_benign, train_malignant = 1032, 2468
    val_benign, val_malignant = 125, 375
    test_benign, test_malignant = 269, 731

    best_epoch = int(summary["best_epoch"])
    best_val_acc = float(summary["best_val_acc"])
    best_val_macro_precision = float(summary["best_val_macro_precision"])
    best_val_macro_recall = float(summary["best_val_macro_recall"])
    best_val_macro_f1 = float(summary["best_val_macro_f1"])
    best_val_weighted_f1 = float(summary["best_val_weighted_f1"])
    best_val_loss = float(best_row["val_loss"])

    test_loss = float(test_report["loss"])
    test_acc = float(test_report["accuracy"])
    test_macro_precision = float(summary["test_macro_precision"])
    test_macro_recall = float(summary["test_macro_recall"])
    test_macro_f1 = float(summary["test_macro_f1"])
    test_weighted_f1 = float(summary["test_weighted_f1"])

    val_cm = best_val["metrics"]["confusion_matrix"]
    test_cm = test_report["metrics"]["confusion_matrix"]
    val_benign_correct, val_benign_wrong = val_cm[0]
    val_malignant_wrong, val_malignant_correct = val_cm[1]
    test_benign_correct, test_benign_wrong = test_cm[0]
    test_malignant_wrong, test_malignant_correct = test_cm[1]

    test_benign_metrics = test_report["metrics"]["class_metrics"]["benign"]
    test_malignant_metrics = test_report["metrics"]["class_metrics"]["malignant"]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = zh(r"\u5b8b\u4f53")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), zh(r"\u5b8b\u4f53"))
    style.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(zh(r"5.X \u516c\u5f00\u6570\u636e\u96c6\u8bad\u7ec3\u7ed3\u679c\u8865\u5145\u8bf4\u660e"))
    set_run_font(r, zh(r"\u9ed1\u4f53"), 14, True)

    add_text_paragraph(
        doc,
        zh(
            r"\u4e3a\u8fdb\u4e00\u6b65\u9a8c\u8bc1\u672c\u6587\u6784\u5efa\u7684\u7532\u72b6\u817a\u7ed3\u8282\u8d85\u58f0\u5f71\u50cf\u826f\u6076\u6027\u5206\u7c7b\u6a21\u578b\u7684\u6709\u6548\u6027\uff0c\u672c\u6587\u5728\u6574\u7406\u540e\u7684 TN5000 \u516c\u5f00\u6570\u636e\u96c6\u4e0a\u5b8c\u6210\u4e86\u8fc1\u79fb\u5b66\u4e60\u8bad\u7ec3\u4e0e\u6d4b\u8bd5\u3002\u6570\u636e\u96c6\u6309\u7167 train\u3001val \u548c test \u4e09\u4e2a\u5b50\u96c6\u8fdb\u884c\u7ec4\u7ec7\uff0c\u5176\u4e2d\u8bad\u7ec3\u96c6\u3001\u9a8c\u8bc1\u96c6\u548c\u6d4b\u8bd5\u96c6\u5206\u522b\u7528\u4e8e\u6a21\u578b\u53c2\u6570\u5b66\u4e60\u3001\u6a21\u578b\u9009\u62e9\u4e0e\u6700\u7ec8\u6cdb\u5316\u6027\u80fd\u8bc4\u4f30\u3002\u5177\u4f53\u800c\u8a00\uff0c\u8bad\u7ec3\u96c6\u5171\u5305\u542b 3500 \u5f20\u56fe\u50cf\uff0c\u9a8c\u8bc1\u96c6\u5171\u5305\u542b 500 \u5f20\u56fe\u50cf\uff0c\u6d4b\u8bd5\u96c6\u5171\u5305\u542b 1000 \u5f20\u56fe\u50cf\uff0c\u80fd\u591f\u8f83\u597d\u5730\u533a\u5206\u6a21\u578b\u5f00\u53d1\u9636\u6bb5\u4e0e\u6700\u7ec8\u8bc4\u4ef7\u9636\u6bb5\u3002"
        ),
    )
    add_text_paragraph(
        doc,
        zh(
            r"\u4ece\u7c7b\u522b\u5206\u5e03\u6765\u770b\uff0c\u8bad\u7ec3\u96c6\u5305\u542b\u826f\u6027\u6837\u672c 1032 \u5f20\u3001\u6076\u6027\u6837\u672c 2468 \u5f20\uff1b\u9a8c\u8bc1\u96c6\u5305\u542b\u826f\u6027\u6837\u672c 125 \u5f20\u3001\u6076\u6027\u6837\u672c 375 \u5f20\uff1b\u6d4b\u8bd5\u96c6\u5305\u542b\u826f\u6027\u6837\u672c 269 \u5f20\u3001\u6076\u6027\u6837\u672c 731 \u5f20\u3002\u8be5\u6570\u636e\u5212\u5206\u65b9\u5f0f\u4e0e\u672c\u6587\u7cfb\u7edf\u7684\u4e8c\u5206\u7c7b\u76ee\u6807\u4fdd\u6301\u4e00\u81f4\uff0c\u65e2\u4fdd\u7559\u4e86\u7532\u72b6\u817a\u7ed3\u8282\u8d85\u58f0\u56fe\u50cf\u5728\u771f\u5b9e\u573a\u666f\u4e2d\u7684\u7c7b\u522b\u5206\u5e03\u7279\u70b9\uff0c\u4e5f\u4e3a\u540e\u7eed\u5206\u6790\u6a21\u578b\u5728\u4e24\u7c7b\u6837\u672c\u4e0a\u7684\u8bc6\u522b\u80fd\u529b\u5dee\u5f02\u63d0\u4f9b\u4e86\u4f9d\u636e\u3002\u8868 5-X \u7ed9\u51fa\u4e86\u672c\u6b21\u5b9e\u9a8c\u6240\u91c7\u7528\u7684\u6570\u636e\u96c6\u5212\u5206\u60c5\u51b5\u3002"
        ),
    )
    add_text_paragraph(
        doc,
        zh(
            r"\u4e3a\u4e86\u66f4\u76f4\u89c2\u5730\u8bf4\u660e\u5404\u5b50\u96c6\u7684\u6837\u672c\u89c4\u6a21\u53ca\u5176\u7c7b\u522b\u6784\u6210\uff0c\u4e0b\u9762\u5bf9\u6570\u636e\u96c6\u5212\u5206\u60c5\u51b5\u8fdb\u884c\u96c6\u4e2d\u5217\u793a\u3002"
        ),
    )

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = cap.add_run(zh(r"\u8868 5-X  TN5000 \u6570\u636e\u96c6\u5212\u5206\u60c5\u51b5"))
    set_run_font(rr, zh(r"\u5b8b\u4f53"), 11)

    table = doc.add_table(rows=4, cols=4)
    style_table(table)
    headers = [
        zh(r"\u6570\u636e\u5b50\u96c6"),
        zh(r"\u826f\u6027\u6837\u672c\u6570"),
        zh(r"\u6076\u6027\u6837\u672c\u6570"),
        zh(r"\u603b\u6570"),
    ]
    for i, h in enumerate(headers):
        set_table_text(table.rows[0].cells[i], h, bold=True)
    for ridx, row in enumerate(
        [
            [zh(r"\u8bad\u7ec3\u96c6"), str(train_benign), str(train_malignant), str(train_benign + train_malignant)],
            [zh(r"\u9a8c\u8bc1\u96c6"), str(val_benign), str(val_malignant), str(val_benign + val_malignant)],
            [zh(r"\u6d4b\u8bd5\u96c6"), str(test_benign), str(test_malignant), str(test_benign + test_malignant)],
        ],
        start=1,
    ):
        for cidx, val in enumerate(row):
            set_table_text(table.rows[ridx].cells[cidx], val)
    add_text_paragraph(
        doc,
        zh(
            r"\u7531\u8868 5-X \u53ef\u4ee5\u770b\u51fa\uff0c\u6574\u4e2a\u6570\u636e\u96c6\u5728\u4e09\u4e2a\u5b50\u96c6\u4e2d\u5747\u4fdd\u6301\u4e86\u826f\u6027\u4e0e\u6076\u6027\u4e8c\u5206\u7c7b\u7ed3\u6784\uff0c\u8fd9\u79cd\u5212\u5206\u65b9\u5f0f\u4e0d\u4ec5\u4fbf\u4e8e\u6a21\u578b\u5f00\u53d1\u9636\u6bb5\u7684\u5bf9\u6bd4\u8bc4\u4f30\uff0c\u4e5f\u6709\u5229\u4e8e\u540e\u7eed\u5728\u72ec\u7acb\u6d4b\u8bd5\u96c6\u4e0a\u5bf9\u6a21\u578b\u6cdb\u5316\u6027\u80fd\u8fdb\u884c\u68c0\u9a8c\u3002"
        ),
    )

    add_text_paragraph(
        doc,
        zh(
            r"\u5728\u8bad\u7ec3\u914d\u7f6e\u65b9\u9762\uff0c\u672c\u6587\u91c7\u7528\u57fa\u4e8e ImageNet \u9884\u8bad\u7ec3\u6743\u91cd\u521d\u59cb\u5316\u7684 ResNet50 \u4f5c\u4e3a\u9aa8\u5e72\u7f51\u7edc\uff0c\u5c06\u5176\u6700\u540e\u5168\u8fde\u63a5\u5c42\u66ff\u6362\u4e3a\u4e8c\u5206\u7c7b\u8f93\u51fa\u5c42\uff0c\u5e76\u5728\u7532\u72b6\u817a\u7ed3\u8282\u8d85\u58f0\u5f71\u50cf\u6570\u636e\u4e0a\u8fdb\u884c\u5fae\u8c03\u8bad\u7ec3\u3002\u8bad\u7ec3\u8fc7\u7a0b\u4e2d\u6700\u5927\u8bad\u7ec3\u8f6e\u6570\u8bbe\u7f6e\u4e3a 150\uff0c\u6279\u91cf\u5927\u5c0f\u8bbe\u7f6e\u4e3a 16\uff0c\u521d\u59cb\u5b66\u4e60\u7387\u8bbe\u7f6e\u4e3a 1\u00d710^-4\uff0c\u6743\u91cd\u8870\u51cf\u7cfb\u6570\u8bbe\u7f6e\u4e3a 1\u00d710^-4\u3002\u8003\u8651\u5230\u8bad\u7ec3\u96c6\u4e2d\u826f\u6027\u4e0e\u6076\u6027\u6837\u672c\u6570\u91cf\u5b58\u5728\u4e00\u5b9a\u5dee\u5f02\uff0c\u8bad\u7ec3\u811a\u672c\u6839\u636e\u7c7b\u522b\u5206\u5e03\u6784\u5efa\u4e86\u52a0\u6743\u4ea4\u53c9\u71b5\u635f\u5931\u51fd\u6570\uff0c\u4ee5\u51cf\u8f7b\u7c7b\u522b\u4e0d\u5e73\u8861\u5bf9\u6a21\u578b\u5b66\u4e60\u8fc7\u7a0b\u5e26\u6765\u7684\u5f71\u54cd\uff1b\u4f18\u5316\u5668\u91c7\u7528 Adam\uff0c\u5e76\u4ee5\u9a8c\u8bc1\u96c6\u51c6\u786e\u7387\u4f5c\u4e3a\u6700\u4f18\u6a21\u578b\u6743\u91cd\u7684\u4fdd\u5b58\u4f9d\u636e\u3002\u6574\u4e2a\u8bad\u7ec3\u8fc7\u7a0b\u603b\u8017\u65f6\u7ea6 20 h 39 min 12 s\uff0c\u6700\u4f18\u6a21\u578b\u51fa\u73b0\u5728\u7b2c 32 \u8f6e\u3002"
        ),
    )
    add_text_paragraph(
        doc,
        zh(
            r"\u4e3a\u4e86\u4fbf\u4e8e\u4ece\u6574\u4f53\u4e0a\u6bd4\u8f83\u6a21\u578b\u5728\u6700\u4f18\u9a8c\u8bc1\u8f6e\u6b21\u4e0e\u6700\u7ec8\u6d4b\u8bd5\u96c6\u4e0a\u7684\u6027\u80fd\u8868\u73b0\uff0c\u672c\u6587\u5bf9\u4e3b\u8981\u5b9e\u9a8c\u6307\u6807\u8fdb\u884c\u4e86\u7edf\u4e00\u6c47\u603b\u3002"
        ),
    )

    cap2 = doc.add_paragraph()
    cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = cap2.add_run(zh(r"\u8868 5-X  \u516c\u5f00\u6570\u636e\u96c6\u8bad\u7ec3\u4e0e\u6d4b\u8bd5\u4e3b\u8981\u7ed3\u679c"))
    set_run_font(rr, zh(r"\u5b8b\u4f53"), 11)

    table2 = doc.add_table(rows=3, cols=7)
    style_table(table2)
    headers2 = [
        zh(r"\u6570\u636e\u96c6"),
        "Loss",
        zh(r"\u51c6\u786e\u7387/%"),
        zh(r"\u5b8f\u5e73\u5747\u7cbe\u786e\u7387/%"),
        zh(r"\u5b8f\u5e73\u5747\u53ec\u56de\u7387/%"),
        zh(r"\u5b8f\u5e73\u5747F1/%"),
        zh(r"\u52a0\u6743F1/%"),
    ]
    for i, h in enumerate(headers2):
        set_table_text(table2.rows[0].cells[i], h, bold=True)
    rows2 = [
        [
            zh(r"\u9a8c\u8bc1\u96c6\u6700\u4f18\u8f6e\u6b21"),
            f"{best_val_loss:.4f}",
            f"{best_val_acc * 100:.2f}",
            f"{best_val_macro_precision * 100:.2f}",
            f"{best_val_macro_recall * 100:.2f}",
            f"{best_val_macro_f1 * 100:.2f}",
            f"{best_val_weighted_f1 * 100:.2f}",
        ],
        [
            zh(r"\u6d4b\u8bd5\u96c6"),
            f"{test_loss:.4f}",
            f"{test_acc * 100:.2f}",
            f"{test_macro_precision * 100:.2f}",
            f"{test_macro_recall * 100:.2f}",
            f"{test_macro_f1 * 100:.2f}",
            f"{test_weighted_f1 * 100:.2f}",
        ],
    ]
    for ridx, row in enumerate(rows2, start=1):
        for cidx, val in enumerate(row):
            set_table_text(table2.rows[ridx].cells[cidx], val)
    add_text_paragraph(
        doc,
        zh(
            r"\u7531\u8868 5-X \u53ef\u4ee5\u770b\u51fa\uff0c\u6a21\u578b\u5728\u9a8c\u8bc1\u96c6\u4e0a\u53d6\u5f97\u4e86\u8f83\u9ad8\u7684\u8bc6\u522b\u7ed3\u679c\uff0c\u5728\u72ec\u7acb\u6d4b\u8bd5\u96c6\u4e0a\u4ecd\u4fdd\u6301\u4e86\u8f83\u597d\u7684\u51c6\u786e\u7387\u548c F1 \u6307\u6807\uff0c\u8bf4\u660e\u8be5\u6a21\u578b\u5728\u672a\u89c1\u6837\u672c\u4e0a\u5177\u6709\u4e00\u5b9a\u7684\u6cdb\u5316\u80fd\u529b\u3002"
        ),
    )

    add_text_paragraph(
        doc,
        zh(
            rf"\u4ece\u9a8c\u8bc1\u96c6\u7ed3\u679c\u6765\u770b\uff0c\u6a21\u578b\u5728\u7b2c {best_epoch} \u8f6e\u8fbe\u5230\u6700\u4f73\u9a8c\u8bc1\u6027\u80fd\uff0c\u5176\u9a8c\u8bc1\u96c6\u635f\u5931\u4e3a {best_val_loss:.4f}\uff0c\u9a8c\u8bc1\u96c6\u51c6\u786e\u7387\u4e3a {best_val_acc * 100:.2f}%\uff0c\u5b8f\u5e73\u5747\u7cbe\u786e\u7387\u4e3a {best_val_macro_precision * 100:.2f}%\uff0c\u5b8f\u5e73\u5747\u53ec\u56de\u7387\u4e3a {best_val_macro_recall * 100:.2f}%\uff0c\u5b8f\u5e73\u5747 F1 \u503c\u4e3a {best_val_macro_f1 * 100:.2f}%\uff0c\u52a0\u6743 F1 \u503c\u4e3a {best_val_weighted_f1 * 100:.2f}%\u3002\u8fd9\u4e00\u7ed3\u679c\u8868\u660e\uff0c\u7ecf\u8fc7\u8fc1\u79fb\u5b66\u4e60\u5fae\u8c03\u540e\u7684 ResNet50 \u5df2\u80fd\u591f\u8f83\u597d\u5730\u63d0\u53d6\u7532\u72b6\u817a\u7ed3\u8282\u8d85\u58f0\u5f71\u50cf\u4e2d\u7684\u5224\u522b\u6027\u7279\u5f81\uff0c\u5e76\u5728\u9a8c\u8bc1\u96c6\u4e0a\u83b7\u5f97\u8f83\u9ad8\u7684\u5206\u7c7b\u7cbe\u5ea6\u3002\u7531\u8bad\u7ec3\u6307\u6807\u53d8\u5316\u66f2\u7ebf\u53ef\u4ee5\u770b\u51fa\uff0c\u6a21\u578b\u5728\u8bad\u7ec3\u521d\u671f\u6536\u655b\u901f\u5ea6\u8f83\u5feb\uff0c\u9a8c\u8bc1\u96c6\u6027\u80fd\u5728\u524d\u82e5\u5e72\u8f6e\u5185\u6301\u7eed\u63d0\u5347\uff0c\u5e76\u5728\u7b2c 32 \u8f6e\u9644\u8fd1\u8fbe\u5230\u5cf0\u503c\u3002\u56fe 5-X \u5c55\u793a\u4e86\u672c\u6b21\u8bad\u7ec3\u8fc7\u7a0b\u4e2d\u8bad\u7ec3\u635f\u5931\u4e0e\u9a8c\u8bc1\u51c6\u786e\u7387\u7684\u6574\u4f53\u53d8\u5316\u8d8b\u52bf\u3002"
        ),
    )
    doc.add_picture(str(ART_DIR / "training_curves.png"), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, zh(r"\u56fe 5-X  \u8bad\u7ec3\u635f\u5931\u4e0e\u9a8c\u8bc1\u51c6\u786e\u7387\u53d8\u5316\u66f2\u7ebf"))
    add_text_paragraph(
        doc,
        zh(
            r"\u4ece\u56fe 5-X \u53ef\u4ee5\u770b\u51fa\uff0c\u8bad\u7ec3\u635f\u5931\u5728\u524d\u671f\u9636\u6bb5\u4e0b\u964d\u5e45\u5ea6\u8f83\u5927\uff0c\u968f\u7740\u8fed\u4ee3\u8f6e\u6b21\u7684\u589e\u52a0\u9010\u6e10\u8d8b\u4e8e\u5e73\u7a33\uff1b\u9a8c\u8bc1\u51c6\u786e\u7387\u5728\u8bad\u7ec3\u521d\u671f\u5feb\u901f\u63d0\u5347\uff0c\u540e\u671f\u5219\u56f4\u7ed5\u8f83\u9ad8\u6c34\u5e73\u5c0f\u5e45\u6ce2\u52a8\uff0c\u8fd9\u4e0e\u6700\u4f18\u9a8c\u8bc1\u8f6e\u6b21\u7684\u9009\u62e9\u7ed3\u679c\u57fa\u672c\u4e00\u81f4\u3002"
        ),
    )

    add_text_paragraph(
        doc,
        zh(
            rf"\u82e5\u7ed3\u5408\u9a8c\u8bc1\u96c6\u6df7\u6dc6\u77e9\u9635\u8fdb\u4e00\u6b65\u5206\u6790\u53ef\u4ee5\u53d1\u73b0\uff0c\u826f\u6027\u6837\u672c\u4e2d\u6709 {val_benign_correct} \u5f20\u88ab\u6b63\u786e\u8bc6\u522b\uff0c{val_benign_wrong} \u5f20\u88ab\u8bef\u5224\u4e3a\u6076\u6027\uff1b\u6076\u6027\u6837\u672c\u4e2d\u6709 {val_malignant_correct} \u5f20\u88ab\u6b63\u786e\u8bc6\u522b\uff0c{val_malignant_wrong} \u5f20\u88ab\u8bef\u5224\u4e3a\u826f\u6027\u3002\u603b\u4f53\u6765\u770b\uff0c\u6a21\u578b\u5728\u9a8c\u8bc1\u9636\u6bb5\u5bf9\u6076\u6027\u6837\u672c\u7684\u8bc6\u522b\u6548\u679c\u76f8\u5bf9\u66f4\u7a33\u5b9a\uff0c\u540c\u65f6\u5bf9\u826f\u6027\u6837\u672c\u4e5f\u4fdd\u6301\u4e86\u8f83\u9ad8\u7684\u8bc6\u522b\u80fd\u529b\u3002\u56fe 5-X \u7ed9\u51fa\u4e86\u6700\u4f18\u9a8c\u8bc1\u8f6e\u6b21\u5bf9\u5e94\u7684\u6df7\u6dc6\u77e9\u9635\uff0c\u53ef\u7528\u4e8e\u8fdb\u4e00\u6b65\u89c2\u5bdf\u6a21\u578b\u5728\u4e24\u7c7b\u7ed3\u8282\u4e0a\u7684\u5206\u7c7b\u5206\u5e03\u3002"
        ),
    )
    doc.add_picture(str(ART_DIR / "confusion_matrix_best.png"), width=Inches(5.4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, zh(r"\u56fe 5-X  \u6700\u4f18\u9a8c\u8bc1\u8f6e\u6b21\u6df7\u6dc6\u77e9\u9635"))
    add_text_paragraph(
        doc,
        zh(
            r"\u7531\u56fe 5-X \u53ef\u4ee5\u76f4\u89c2\u770b\u51fa\uff0c\u6700\u4f18\u9a8c\u8bc1\u8f6e\u6b21\u4e0b\u6a21\u578b\u5bf9\u6076\u6027\u6837\u672c\u7684\u8bc6\u522b\u66f4\u4e3a\u7a33\u5b9a\uff0c\u540c\u65f6\u826f\u6027\u6837\u672c\u4ecd\u4fdd\u6301\u4e86\u8f83\u9ad8\u7684\u8bc6\u522b\u6c34\u5e73\uff0c\u8bf4\u660e\u6a21\u578b\u5df2\u80fd\u8f83\u597d\u5730\u5b66\u4e60\u4e24\u7c7b\u56fe\u50cf\u7684\u533a\u5206\u7279\u5f81\u3002"
        ),
    )

    add_text_paragraph(
        doc,
        zh(
            rf"\u5728\u72ec\u7acb\u6d4b\u8bd5\u96c6\u4e0a\u7684\u8bc4\u4f30\u7ed3\u679c\u663e\u793a\uff0c\u6a21\u578b\u6d4b\u8bd5\u635f\u5931\u4e3a {test_loss:.4f}\uff0c\u6d4b\u8bd5\u51c6\u786e\u7387\u4e3a {test_acc * 100:.2f}%\uff0c\u5b8f\u5e73\u5747\u7cbe\u786e\u7387\u4e3a {test_macro_precision * 100:.2f}%\uff0c\u5b8f\u5e73\u5747\u53ec\u56de\u7387\u4e3a {test_macro_recall * 100:.2f}%\uff0c\u5b8f\u5e73\u5747 F1 \u503c\u4e3a {test_macro_f1 * 100:.2f}%\uff0c\u52a0\u6743 F1 \u503c\u4e3a {test_weighted_f1 * 100:.2f}%\u3002\u4ece\u6d4b\u8bd5\u96c6\u6df7\u6dc6\u77e9\u9635\u6765\u770b\uff0c{test_benign} \u5f20\u826f\u6027\u6837\u672c\u4e2d\u6709 {test_benign_correct} \u5f20\u88ab\u6b63\u786e\u8bc6\u522b\uff0c{test_benign_wrong} \u5f20\u88ab\u8bef\u5224\u4e3a\u6076\u6027\uff1b{test_malignant} \u5f20\u6076\u6027\u6837\u672c\u4e2d\u6709 {test_malignant_correct} \u5f20\u88ab\u6b63\u786e\u8bc6\u522b\uff0c{test_malignant_wrong} \u5f20\u88ab\u8bef\u5224\u4e3a\u826f\u6027\u3002\u6309\u7c7b\u522b\u5206\u6790\uff0c\u6a21\u578b\u5bf9\u6076\u6027\u6837\u672c\u7684\u7cbe\u786e\u7387\u8fbe\u5230 {test_malignant_metrics['precision'] * 100:.2f}%\uff0c\u53ec\u56de\u7387\u8fbe\u5230 {test_malignant_metrics['recall'] * 100:.2f}%\uff0c\u5bf9\u5e94 F1 \u503c\u4e3a {test_malignant_metrics['f1'] * 100:.2f}%\uff1b\u76f8\u6bd4\u4e4b\u4e0b\uff0c\u826f\u6027\u6837\u672c\u7684\u7cbe\u786e\u7387\u4e3a {test_benign_metrics['precision'] * 100:.2f}%\uff0c\u53ec\u56de\u7387\u4e3a {test_benign_metrics['recall'] * 100:.2f}%\uff0c\u5bf9\u5e94 F1 \u503c\u4e3a {test_benign_metrics['f1'] * 100:.2f}%\u3002\u8fd9\u4e00\u73b0\u8c61\u8bf4\u660e\uff0c\u5728\u5f53\u524d\u6570\u636e\u5206\u5e03\u548c\u8bad\u7ec3\u6761\u4ef6\u4e0b\uff0c\u6a21\u578b\u5bf9\u6076\u6027\u7279\u5f81\u7684\u5b66\u4e60\u66f4\u52a0\u5145\u5206\uff0c\u800c\u5bf9\u826f\u6027\u6837\u672c\u7684\u8bc6\u522b\u4ecd\u6709\u8fdb\u4e00\u6b65\u63d0\u5347\u7a7a\u95f4\u3002"
        ),
    )
    doc.add_picture(str(ART_DIR / "confusion_matrix_test.png"), width=Inches(5.4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, zh(r"\u56fe 5-X  \u6d4b\u8bd5\u96c6\u6df7\u6dc6\u77e9\u9635"))
    add_text_paragraph(
        doc,
        zh(
            r"\u4ece\u56fe 5-X \u6240\u793a\u7684\u6d4b\u8bd5\u96c6\u6df7\u6dc6\u77e9\u9635\u53ef\u4ee5\u8fdb\u4e00\u6b65\u770b\u51fa\uff0c\u6a21\u578b\u5728\u72ec\u7acb\u6d4b\u8bd5\u6570\u636e\u4e0a\u4ecd\u80fd\u591f\u5bf9\u5927\u90e8\u5206\u6076\u6027\u6837\u672c\u7ed9\u51fa\u6b63\u786e\u5224\u65ad\uff0c\u4f46\u5bf9\u90e8\u5206\u826f\u6027\u56fe\u50cf\u4ecd\u5b58\u5728\u4e00\u5b9a\u8bef\u5224\u60c5\u51b5\uff0c\u8fd9\u4e5f\u4e0e\u524d\u6587\u5bf9\u7c7b\u522b\u7ea7\u6307\u6807\u7684\u5206\u6790\u7ed3\u8bba\u76f8\u4e92\u5370\u8bc1\u3002"
        ),
    )

    add_text_paragraph(
        doc,
        zh(
            rf"\u4ece\u8bad\u7ec3\u540e\u671f\u7684\u6307\u6807\u53d8\u5316\u8d8b\u52bf\u6765\u770b\uff0c\u6a21\u578b\u5728\u7b2c {best_epoch} \u8f6e\u4e4b\u540e\u8bad\u7ec3\u96c6\u51c6\u786e\u7387\u4ecd\u7ee7\u7eed\u4e0a\u5347\uff0c\u800c\u9a8c\u8bc1\u96c6\u51c6\u786e\u7387\u53ca\u5b8f\u5e73\u5747 F1 \u503c\u5728\u540e\u7eed\u8f6e\u6b21\u4e2d\u51fa\u73b0\u4e00\u5b9a\u6ce2\u52a8\uff0c\u8bf4\u660e\u6a21\u578b\u5728\u8bad\u7ec3\u540e\u671f\u5b58\u5728\u8f7b\u5fae\u8fc7\u62df\u5408\u8d8b\u52bf\u3002\u56e0\u6b64\uff0c\u672c\u6587\u5728\u5b9e\u9a8c\u5b9e\u73b0\u4e2d\u91c7\u7528\u201c\u4fdd\u5b58\u9a8c\u8bc1\u96c6\u6700\u4f18\u6743\u91cd\u201d\u7684\u7b56\u7565\uff0c\u5c06\u7b2c {best_epoch} \u8f6e\u5bf9\u5e94\u7684\u6a21\u578b\u53c2\u6570\u4f5c\u4e3a\u6700\u7ec8\u63a8\u7406\u6a21\u578b\uff0c\u800c\u4e0d\u662f\u76f4\u63a5\u91c7\u7528\u6700\u540e\u4e00\u8f6e\u8bad\u7ec3\u53c2\u6570\u3002\u8be5\u7b56\u7565\u6709\u52a9\u4e8e\u5728\u4fdd\u8bc1\u6a21\u578b\u5145\u5206\u5b66\u4e60\u7684\u540c\u65f6\uff0c\u964d\u4f4e\u540e\u671f\u8fc7\u62df\u5408\u5bf9\u7cfb\u7edf\u5b9e\u9645\u9884\u6d4b\u6027\u80fd\u7684\u5f71\u54cd\u3002"
        ),
    )
    add_text_paragraph(
        doc,
        zh(
            r"\u603b\u4f53\u6765\u770b\uff0c\u57fa\u4e8e TN5000 \u516c\u5f00\u6570\u636e\u96c6\u7684\u5b9e\u9a8c\u7ed3\u679c\u8868\u660e\uff0c\u672c\u6587\u91c7\u7528\u7684\u8fc1\u79fb\u5b66\u4e60\u5206\u7c7b\u65b9\u6848\u80fd\u591f\u5728\u7532\u72b6\u817a\u7ed3\u8282\u8d85\u58f0\u5f71\u50cf\u826f\u6076\u6027\u4e8c\u5206\u7c7b\u4efb\u52a1\u4e2d\u53d6\u5f97\u8f83\u597d\u7684\u6548\u679c\uff0c\u5e76\u5177\u5907\u4e00\u5b9a\u7684\u6cdb\u5316\u80fd\u529b\u3002\u8be5\u7ed3\u679c\u4e3a\u540e\u7eed\u7cfb\u7edf\u4e2d\u7684\u6a21\u578b\u63a8\u7406\u6a21\u5757\u63d0\u4f9b\u4e86\u5b9e\u9a8c\u4f9d\u636e\uff0c\u4e5f\u8bf4\u660e\u5c06\u9884\u8bad\u7ec3 ResNet50 \u4e0e Grad-CAM \u53ef\u89e3\u91ca\u5316\u65b9\u6cd5\u7ed3\u5408\u5e94\u7528\u4e8e\u7532\u72b6\u817a\u7ed3\u8282\u8d85\u58f0\u8f85\u52a9\u5206\u6790\u7cfb\u7edf\u5177\u6709\u5b9e\u9645\u53ef\u884c\u6027\u3002"
        ),
    )

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
